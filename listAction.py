#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Created on Sat Jul ” 7 06:00:41 2018

@author: fsakai
"""
file = "./20180707_Sat.docx"
marklist = ["★", "☆", ]


def displayMarkedSentenceFromDocument(file, marklist):
    from docx import Document
    doc = Document(file)
    for mark in marklist:
        print("displaying sentence marked with: " + mark)
        for para in doc.paragraphs:
            text = para.text
            search_from, mark_from = 0, 0
            while mark_from >= 0:
                mark_from = text.find(mark, search_from)
                search_from = mark_from + 1
                if mark_from != -1:
                    mark_end = text.find("\\"+mark, mark_from)
                    if mark_end == -1:
                        mark_end = len(text)
                    else: 
                        search_from = mark_end + 2
                    print(text[mark_from:mark_end], )

                    
